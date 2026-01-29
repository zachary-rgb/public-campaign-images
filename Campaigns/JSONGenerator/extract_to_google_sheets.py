"""
Walgreens Campaign Content Extractor for Google Sheets
Extracts content from Word documents and prepares for direct paste into Google Sheets
"""

from docx import Document
import json
import csv
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional

try:
    import pyperclip
    CLIPBOARD_AVAILABLE = True
except ImportError:
    CLIPBOARD_AVAILABLE = False


class CampaignExtractor:
    """Extract campaign content from Walgreens Word documents"""
    
    # Standard column order matching your Google Sheet
    STANDARD_COLUMNS = [
        # Note: "Generate JSON" and "Campaign ID" columns are managed by the spreadsheet, not extracted
        'Campaign Name',
        'Message Name',
        'Email Subject Line',
        'Language',
        'Hero Image (URL)',
        'Banner Headline',
        'Study Information',
        'CTA 1',
        'Eligibility',
        'What to Expect',
        'CTA 2',
        'Closing',
        'Optional Resource',
        'End Matter (Enter Sponsor)',
        'URL/UTM'
    ]
    
    # Enable color and checkbox detection
    USE_COLOR_DETECTION = True
    USE_CHECKBOX_DETECTION = True
    
    # Map Word doc section labels to spreadsheet columns
    COLUMN_MAPPING = {
        'EMAIL SUBJECT LINE OPTIONS': 'Email Subject Line',
        'EMAIL SUBJECT LINE': 'Email Subject Line',
        'SUBJECT LINE': 'Email Subject Line',
        'SALUTATION': 'Salutation',
        'BANNER HEADLINE OPTIONS': 'Banner Headline',
        'BANNER HEADLINE': 'Banner Headline',
        'HEADLINE': 'Banner Headline',
        'STUDY INFORMATION': 'Study Information',
        'REMINDER EMAIL COPY': 'Study Information',  # Alternative label for same field
        'CTA 1': 'CTA 1',
        'CTA 2': 'CTA 2',
        'BUTTON 1': 'CTA 1',
        'BUTTON 2': 'CTA 2',
        'ELIGIBILITY': 'Eligibility',
        'WHAT TO EXPECT': 'What to Expect',
        'CLOSING': 'Closing',
        'OPTIONAL RESOURCE': 'Optional Resource',
        'RESOURCE': 'Optional Resource',
        'HERO SPACE IMAGE OPTIONS': 'Hero Image (URL)',
        'HERO IMAGE': 'Hero Image (URL)',
        'IMAGE': 'Hero Image (URL)',
        'LOGO OPTIONS': 'Logo',
        'LOGO': 'Logo'
    }
    
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)
        self.warnings = []  # Track warnings about missing selections
        self.warned_fields = set()  # Track which fields we've already warned about
        self.current_message_name = None  # Track current message being processed
        
    def is_green_text(self, run):
        """Check if text is green (variable content in Walgreens template)"""
        try:
            if not run.font.color or not run.font.color.rgb:
                return False
            rgb = run.font.color.rgb
            
            # RGBColor in python-docx is an integer, extract components
            if isinstance(rgb, int):
                r = (rgb >> 16) & 0xFF
                g = (rgb >> 8) & 0xFF
                b = rgb & 0xFF
            else:
                # Handle tuple-like objects
                r, g, b = rgb[0], rgb[1], rgb[2]
            
            # Green color detection - adjust thresholds as needed
            # Typical green: high green value, low red and blue
            return (r < 100 and g > 150 and b < 100)
        except:
            return False
    
    def is_red_text(self, run):
        """Check if text is red (Walgreens standard content - not editable)"""
        try:
            if not run.font.color or not run.font.color.rgb:
                return False
            rgb = run.font.color.rgb
            
            # RGBColor in python-docx is an integer, extract components
            if isinstance(rgb, int):
                r = (rgb >> 16) & 0xFF
                g = (rgb >> 8) & 0xFF
                b = rgb & 0xFF
            else:
                # Handle tuple-like objects
                r, g, b = rgb[0], rgb[1], rgb[2]
            
            # Red color detection
            return (r > 150 and g < 100 and b < 100)
        except:
            return False
    
    def get_text_color_type(self, cell):
        """
        Determine if cell contains GREEN (variable), RED (standard), or neutral text
        Returns: 'green', 'red', or 'neutral'
        """
        if not self.USE_COLOR_DETECTION:
            return 'neutral'
        
        green_chars = 0
        red_chars = 0
        total_chars = 0
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                text_len = len(run.text.strip())
                if text_len == 0:
                    continue
                total_chars += text_len
                
                if self.is_green_text(run):
                    green_chars += text_len
                elif self.is_red_text(run):
                    red_chars += text_len
        
        if total_chars == 0:
            return 'neutral'
        
        # If majority is green (>50%)
        if green_chars > total_chars * 0.5:
            return 'green'
        # If majority is red (>50%)
        elif red_chars > total_chars * 0.5:
            return 'red'
        else:
            return 'neutral'
    
    def has_multiple_options(self, text: str) -> bool:
        """Check if text contains multiple numbered/lettered options"""
        lines = text.split('\n')
        option_count = 0
        for line in lines:
            # Check for numbered or lettered options
            if re.match(r'^\s*\d+\.\s+', line) or re.match(r'^\s*[A-Z]\.\s+', line, re.IGNORECASE):
                option_count += 1
        return option_count > 1
    def combine_salutation_with_study_info(self, data: Dict[str, str]) -> Dict[str, str]:
        """
        Combine SALUTATION with STUDY INFORMATION/REMINDER EMAIL COPY.
        - If SALUTATION is RED (or missing) → APPEND default: "**Dear [Patient First Name]**"
        - If SALUTATION is GREEN → APPEND actual content from document
        Always prepends to STUDY INFORMATION content.
        """
        salutation = data.get('SALUTATION', '').strip()
        salutation_is_red = data.get('SALUTATION_IS_RED', False)  # Track if field was red
        study_info = data.get('STUDY INFORMATION', '').strip()
        
        # If STUDY INFORMATION is empty but we have REMINDER EMAIL COPY, use that
        if not study_info:
            study_info = data.get('REMINDER EMAIL COPY', '').strip()
        
        # Determine salutation to use
        if salutation_is_red or not salutation:
            # RED or missing → Use Walgreens standard default
            final_salutation = "**Dear [Patient First Name]**"
        else:
            # GREEN → Use actual content from document
            final_salutation = salutation
        
        # ALWAYS append salutation to study info (prepend salutation to content)
        if study_info:
            # Append default/custom salutation before the study info content
            data['STUDY INFORMATION'] = final_salutation + '\n\n' + study_info
        else:
            # No study info content, just use salutation
            data['STUDY INFORMATION'] = final_salutation
        
        # Remove standalone SALUTATION field
        if 'SALUTATION' in data:
            del data['SALUTATION']
        if 'SALUTATION_IS_RED' in data:
            del data['SALUTATION_IS_RED']
        
        return data

    

    def apply_red_field_defaults(self, data: Dict[str, str]) -> Dict[str, str]:
        """
        Apply Walgreens standard defaults for RED fields.
        - CLOSING: If RED or missing → "Sincerely,\nWalgreens Clinical Trials"
        - CLOSING: If GREEN → Use actual content
        """
        closing = data.get('CLOSING', '').strip()
        closing_is_red = data.get('CLOSING_IS_RED', False)
        
        # Handle CLOSING field
        if closing_is_red or not closing:
            # RED or missing → Use Walgreens standard default
            data['CLOSING'] = "Sincerely,\nWalgreens Clinical Trials"
        # else: GREEN → keep actual content (already in data['CLOSING'])
        
        # Clean up tracking fields
        if 'CLOSING_IS_RED' in data:
            del data['CLOSING_IS_RED']
        
        return data

    
    def extract_cta_from_study_info(self, data: Dict[str, str]) -> Dict[str, str]:
        """
        Extract CTA and Closing from STUDY INFORMATION field
        Pattern: "Learn more [INSERT CTA BUTTON]" → CTA 1
        Then "CTA" labeled field → CTA 2
        Text below initial CTA → Closing
        """
        study_info = data.get('STUDY INFORMATION', '').strip()
        
        if not study_info:
            return data
        
        # Check for "Learn more [INSERT CTA BUTTON]" pattern
        learn_more_pattern = r'Learn more\s*\[INSERT CTA BUTTON\]'
        match = re.search(learn_more_pattern, study_info, re.IGNORECASE)
        
        if match:
            # Found the pattern
            # Split content at this point
            before_cta = study_info[:match.start()].strip()
            cta_text = study_info[match.start():match.end()].strip()
            after_cta = study_info[match.end():].strip()
            
            # Update Study Information (content before CTA)
            data['STUDY INFORMATION'] = before_cta
            
            # Extract CTA 1 (the Learn more button)
            if not data.get('CTA 1'):  # Only if CTA 1 is empty
                data['CTA 1'] = cta_text
            
            # Parse content after the Learn more CTA
            if after_cta:
                # Look for another CTA label in the remaining text
                # Pattern: lines starting with "CTA" or containing CTA-like text
                lines = after_cta.split('\n')
                cta2_found = False
                cta2_lines = []
                closing_lines = []
                
                for line in lines:
                    if not cta2_found and ('CTA' in line.upper() or 'BUTTON' in line.upper()):
                        cta2_found = True
                        cta2_lines.append(line)
                    elif cta2_found and line.strip() and not any(keyword in line.upper() for keyword in ['CLOSING', 'REGARDS', 'SINCERELY']):
                        # Still part of CTA 2
                        cta2_lines.append(line)
                    elif line.strip():
                        # This is closing content
                        closing_lines.append(line)
                
                # Set CTA 2 if found
                if cta2_lines and not data.get('CTA 2'):
                    data['CTA 2'] = '\n'.join(cta2_lines).strip()
                
                # Set Closing if found
                if closing_lines and not data.get('CLOSING'):
                    data['CLOSING'] = '\n'.join(closing_lines).strip()
                elif not closing_lines and not cta2_found:
                    # All after_cta content goes to closing
                    if not data.get('CLOSING'):
                        data['CLOSING'] = after_cta
        
        return data

    def extract_selected_option(self, text: str, field_label: str = None) -> str:
        """
        Extract option marked with checkbox.
        Supports: [X], [x], ☑ (U+2611 checked), ☐ (U+2610 unchecked), [✓], ***
        Falls back to first option if no checkbox found.
        Logs warning if multi-option field has no checkbox.
        """
        if not self.USE_CHECKBOX_DETECTION:
            return self.extract_first_option(text)
        
        lines = text.split('\n')
        selected_lines = []
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Check for CHECKED checkbox formats (not unchecked ☐)
            # Must explicitly check for checked boxes, not unchecked
            has_checked = any(marker in line_stripped for marker in 
                             ['[X]', '[x]', '[✓]', '[√]', '☑', '✓', '✔', '[✔]'])
            
            # Skip lines that have unchecked boxes
            has_unchecked = '☐' in line_stripped or '[ ]' in line_stripped
            
            if has_checked and not (has_unchecked and not has_checked):
                # Remove checkbox markers (both checked and unchecked for cleaning)
                cleaned = line_stripped
                # Remove all checkbox marker variations
                for marker in ['[X]', '[x]', '[✓]', '[√]', '☑', '☐', '✓', '✔', '[✔]', '[ ]', '[]', '***']:
                    cleaned = cleaned.replace(marker, '')
                # Remove leading numbers (1., 2., etc.) and letters (A., B., etc.)
                cleaned = re.sub(r'^\s*\d+\.\s*', '', cleaned)
                cleaned = re.sub(r'^\s*[A-Z]\.\s*', '', cleaned, flags=re.IGNORECASE)
                cleaned = cleaned.strip()
                if cleaned:
                    selected_lines.append(cleaned)
        
        # If we found checked items, return them
        if selected_lines:
            return '\n'.join(selected_lines)
        
        # Check if this is a multi-option field without selection
        if self.has_multiple_options(text):
            # Create a unique key: message_name + field_name
            field_name = field_label or 'Unknown Field'
            message_name = self.current_message_name or 'Unknown Email'
            warning_key = f"{message_name}::{field_name}"
            
            # Only add warning if we haven't warned about this specific field in this specific message
            if warning_key not in self.warned_fields:
                self.warned_fields.add(warning_key)
                warning_msg = f"No checkbox found for '{field_name}' (has multiple options)"
                self.warnings.append({
                    'message_name': message_name,
                    'field': field_name,
                    'message': warning_msg,
                    'action': 'Using first option as default'
                })
        
        # Otherwise, extract first option (original behavior)
        return self.extract_first_option(text)
    
    def extract_smart_defaults_from_document(self) -> Dict[str, str]:
        """
        Extract smart default values from document filename and content.
        Returns dict with Campaign Name, Sponsor, Language, URL/UTM, etc.
        """
        defaults = {}
        
        # 1. Extract Campaign Name from filename
        # Pattern: "ClientName_CampaignType_Details.docx" → "ClientName CampaignType Details"
        filename = self.docx_path.stem  # Get filename without extension
        
        # Clean up filename: replace underscores/hyphens with spaces, title case
        campaign_name_from_file = filename.replace('_', ' ').replace('-', ' ')
        # Remove common suffixes like "template", "draft", "v1", etc.
        campaign_name_from_file = re.sub(r'\s+(template|draft|final|v\d+|copy)(\s|$)', '', campaign_name_from_file, flags=re.IGNORECASE)
        campaign_name_from_file = campaign_name_from_file.strip()
        
        if campaign_name_from_file:
            defaults['Campaign Name'] = campaign_name_from_file
        
        # 2. Extract Sponsor from document content
        # Look in common places: End Matter, footer, headers
        sponsor_keywords = ['Takeda', 'Pfizer', 'Novartis', 'Merck', 'AbbVie', 'Sponsored by', 'Brought to you by']
        
        # Search in paragraphs for sponsor mentions
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            # Look for "Sponsored by X" or company names
            for keyword in sponsor_keywords:
                if keyword in text:
                    # Try to extract sponsor name
                    if 'sponsored by' in text.lower() or 'brought to you by' in text.lower():
                        # Extract text after these phrases
                        match = re.search(r'(?:sponsored by|brought to you by)\s+([A-Z][a-zA-Z\s&]+)', text, re.IGNORECASE)
                        if match:
                            defaults['Sponsor'] = match.group(1).strip()
                            break
                    else:
                        # Just found a company name
                        defaults['Sponsor'] = keyword
                        break
            if 'Sponsor' in defaults:
                break
        
        # 3. Detect Language from filename or content
        # Check for language indicators in filename
        filename_lower = filename.lower()
        if 'spanish' in filename_lower or '_es' in filename_lower or '-es' in filename_lower:
            defaults['Language'] = 'en-es'
        elif 'english' in filename_lower or '_en' in filename_lower or '-en' in filename_lower:
            defaults['Language'] = 'en-us'
        else:
            # Default to en-us
            defaults['Language'] = 'en-us'
        
        # 4. Extract URL/UTM from document content
        # Look for URLs in the document
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        for paragraph in self.doc.paragraphs:
            urls = re.findall(url_pattern, paragraph.text)
            if urls:
                # Use the first URL found (often the main campaign URL)
                defaults['URL/UTM'] = urls[0]
                break
        
        # Also check in table cells for URLs
        if 'URL/UTM' not in defaults:
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        urls = re.findall(url_pattern, cell.text)
                        if urls:
                            defaults['URL/UTM'] = urls[0]
                            break
                    if 'URL/UTM' in defaults:
                        break
                if 'URL/UTM' in defaults:
                    break
        
        return defaults
    
    def detect_images_in_document(self):
        """
        Detect all images in the document and extract their information.
        Also checks for [X] markers to identify "preferred" images.
        Returns: List of dicts with image metadata
        """
        images = []
        image_counter = 1
        
        # First, collect all image relationships
        image_rels = []
        try:
            for rel in self.doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_rels.append(rel)
        except Exception as e:
            print(f"Warning: Error detecting images: {e}")
            return images
        
        # Now find images in the document content and check for [X] markers
        image_locations = []  # Track where images appear and if they have [X]
        
        # Check paragraphs for images
        for para in self.doc.paragraphs:
            # Check if paragraph contains an image
            para_has_image = False
            for run in para.runs:
                if run._element.xpath('.//pic:pic'):
                    para_has_image = True
                    break
            
            if para_has_image:
                # Check if paragraph text contains [X] marker
                para_text = para.text.strip()
                has_checkbox = self.has_checkbox_marker(para_text)
                image_locations.append({
                    'has_checkbox': has_checkbox,
                    'context': para_text[:50] if para_text else ''
                })
        
        # Check table cells for images
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Check if cell contains an image
                    cell_has_image = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run._element.xpath('.//pic:pic'):
                                cell_has_image = True
                                break
                        if cell_has_image:
                            break
                    
                    if cell_has_image:
                        # Check if cell text contains [X] marker
                        cell_text = cell.text.strip()
                        has_checkbox = self.has_checkbox_marker(cell_text)
                        image_locations.append({
                            'has_checkbox': has_checkbox,
                            'context': cell_text[:50] if cell_text else ''
                        })
        
        # Build image list with preference detection
        for idx, rel in enumerate(image_rels):
            try:
                image_data = rel.target_part.blob
                
                # Determine image format from content type
                content_type = rel.target_part.content_type
                ext = content_type.split('/')[-1]
                if ext == 'jpeg':
                    ext = 'jpg'
                
                # Get image size (in bytes)
                size_kb = len(image_data) / 1024
                
                # Check if this image is marked as preferred
                is_preferred = False
                context = ''
                if idx < len(image_locations):
                    is_preferred = image_locations[idx]['has_checkbox']
                    context = image_locations[idx]['context']
                
                images.append({
                    'id': image_counter,
                    'filename': f'image_{image_counter}.{ext}',
                    'format': ext.upper(),
                    'size_kb': round(size_kb, 2),
                    'data': image_data,
                    'selected': is_preferred,  # Only select if marked with [X]
                    'preferred': is_preferred,  # Track preferred status
                    'context': context  # Context text for display
                })
                
                image_counter += 1
            except Exception as e:
                print(f"Warning: Could not extract image {image_counter}: {e}")
                continue
        
        return images
    
    def has_checkbox_marker(self, text):
        """
        Check if text contains a checkbox marker indicating selection.
        Same logic as extract_selected_option method.
        """
        if not text:
            return False
        
        # Check for various checkbox markers
        checkbox_markers = ['[X]', '[x]', '[✓]', '[√]', '[✔]', '☑', '✓', '✔']
        text_upper = text.upper()
        
        for marker in checkbox_markers:
            if marker.upper() in text_upper or marker in text:
                return True
        
        return False
    
    def save_images_to_folder(self, images, output_folder='campaign_images'):
        """
        Save detected images to a local folder
        Returns: List of saved file paths
        """
        output_path = Path(output_folder)
        output_path.mkdir(exist_ok=True)
        
        saved_files = []
        for img in images:
            if img.get('selected', False):
                filepath = output_path / img['filename']
                with open(filepath, 'wb') as f:
                    f.write(img['data'])
                saved_files.append(str(filepath))
        
        return saved_files
    
    def extract_email_headers_from_doc(self) -> Dict[int, str]:
        """
        Extract email headers like "Email 1: Long-form email" that appear above tables.
        Returns a dict mapping table index to header text.
        """
        email_headers = {}
        
        try:
            # Find all paragraphs that match "Email X: description" pattern
            for para_idx, paragraph in enumerate(self.doc.paragraphs):
                text = paragraph.text.strip()
                
                # Look for "Email 1:", "Email 2:", etc.
                import re
                match = re.match(r'Email\s+(\d+):\s*(.+)', text, re.IGNORECASE)
                if match:
                    email_num = int(match.group(1))
                    full_text = text  # "Email 1: Long-form email"
                    email_headers[email_num] = full_text
            
            return email_headers
        except:
            return {}
    
    def get_campaign_name_for_table(self, table_index: int) -> str:
        """
        Get the campaign name (email header) for a specific table.
        Table index is 0-based.
        """
        email_headers = self.extract_email_headers_from_doc()
        
        # Email numbers are 1-based, table index is 0-based
        # So table 0 = Email 1, table 1 = Email 2, etc.
        email_num = table_index + 1
        
        return email_headers.get(email_num, f"Email {email_num}")
    
    def extract_table_content(self) -> List[Dict[str, any]]:
        """
        Extract all content from Word doc tables with color and checkbox detection.
        Returns list of dicts, each containing table_index and extracted data.
        """
        all_tables_data = []
        
        for table_idx, table in enumerate(self.doc.tables):
            extracted_data = {}
            
            # Track which table this data came from
            table_campaign_name = self.get_campaign_name_for_table(table_idx)
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    label_cell = cells[0]
                    content_cell = cells[1]
                    
                    # Extract label, handling multi-line text and parenthetical notes
                    label_text = label_cell.text.strip()
                    if not label_text:
                        continue
                    
                    # Get only the first line (before any newlines)
                    label = label_text.split('\n')[0].strip()
                    
                    # Remove any parenthetical notes like "(Email pre read—recommended max: 90 characters)"
                    label = re.sub(r'\s*\([^)]*\).*$', '', label).strip()
                    
                    if not label:
                        continue
                    
                    # Check color of label cell
                    label_color = self.get_text_color_type(label_cell)
                    
                    # Handle RED sections (Walgreens standard)
                    if label_color == 'red':
                        # Special handling for SALUTATION and CLOSING - track but don't extract content
                        label_upper = label.upper()
                        if 'SALUTATION' in label_upper:
                            extracted_data['SALUTATION_IS_RED'] = True
                            extracted_data['SALUTATION'] = ''  # Will use default
                            print(f"  [RED/Standard] {label} → Will use Walgreens default")
                        elif 'CLOSING' in label_upper:
                            extracted_data['CLOSING_IS_RED'] = True
                            extracted_data['CLOSING'] = ''  # Will use default
                            print(f"  [RED/Standard] {label} → Will use Walgreens default")
                        else:
                            # Other RED fields - skip entirely
                            print(f"  [Skipping RED/Standard] {label}")
                        continue
                    
                    content = content_cell.text.strip()
                    if not content:
                        continue
                    
                    # For GREEN sections (variable content), look for checkboxes
                    if label_color == 'green' and self.USE_CHECKBOX_DETECTION:
                        content = self.extract_selected_option(content, label)
                    
                    extracted_data[label] = content
            
            # Add this table's data with its campaign name
            if extracted_data:  # Only add if we extracted something
                all_tables_data.append({
                    'campaign_name': table_campaign_name,
                    'data': extracted_data
                })
        
        return all_tables_data
    
    def normalize_label(self, label: str) -> str:
        """Normalize section labels for matching"""
        normalized = label.upper()
        # Remove content in parentheses
        normalized = re.sub(r'\([^)]*\)', '', normalized)
        # Remove "OPTIONS" suffix variations
        normalized = re.sub(r'\s+OPTIONS?$', '', normalized)
        # Remove common descriptors
        normalized = re.sub(r'\s+(email body copy|recommended max.*)', '', normalized, flags=re.IGNORECASE)
        return normalized.strip()
    
    def map_to_columns(self, extracted_data: Dict[str, str]) -> Dict[str, str]:
        """Map extracted content to spreadsheet columns"""
        spreadsheet_row = {}
        
        for doc_label, content in extracted_data.items():
            normalized_label = self.normalize_label(doc_label)
            
            # Try exact match first
            if normalized_label in self.COLUMN_MAPPING:
                column_name = self.COLUMN_MAPPING[normalized_label]
                spreadsheet_row[column_name] = content
                continue
            
            # Try partial match
            for key, column_name in self.COLUMN_MAPPING.items():
                if key in normalized_label or normalized_label in key:
                    spreadsheet_row[column_name] = content
                    break
        
        return spreadsheet_row
    
    def extract_options(self, text: str) -> List[str]:
        """Extract multiple options from numbered/bulleted text"""
        if not text:
            return []
        
        options = []
        
        # Try numbered list (1. , 2. , 3. etc.)
        numbered = re.split(r'\n\s*\d+\.\s+', text)
        if len(numbered) > 1:
            # First element is text before first number
            options = [opt.strip() for opt in numbered[1:] if opt.strip()]
            if options:
                return options
        
        # Try lettered list (A. , B. , C. etc.)
        lettered = re.split(r'\n\s*[A-Z]\.\s+', text, flags=re.IGNORECASE)
        if len(lettered) > 1:
            options = [opt.strip() for opt in lettered[1:] if opt.strip()]
            if options:
                return options
        
        # Try bullet points
        bulleted = re.split(r'\n\s*[-•●]\s+', text)
        if len(bulleted) > 1:
            options = [opt.strip() for opt in bulleted if opt.strip()]
            if options:
                return options
        
        # Return as single option
        return [text.strip()]
    
    def extract_first_option(self, text: str) -> str:
        """Get the first option from multi-option text"""
        options = self.extract_options(text)
        result = options[0] if options else ''
        
        # Clean up empty checkbox markers [ ] that might be in the text
        result = result.replace('[ ]', '').replace('[]', '')
        # Also remove leading/trailing whitespace from cleanup
        result = result.strip()
        
        return result

    def extract_text_with_formatting(self, cell) -> str:
        """
        Extract text from cell with markdown formatting preserved.
        Converts Word formatting to Markdown:
        - Bold → **text**
        - Italic → *text*
        - Lists → - item or 1. item
        """
        markdown_parts = []
        
        for paragraph in cell.paragraphs:
            para_parts = []
            
            for run in paragraph.runs:
                text = run.text
                if not text:
                    continue
                
                # Apply markdown formatting based on Word styles
                if run.bold and run.italic:
                    text = f"***{text}***"
                elif run.bold:
                    text = f"**{text}**"
                elif run.italic:
                    text = f"*{text}*"
                
                # Handle underline as emphasis (optional)
                if run.underline and not run.bold and not run.italic:
                    text = f"_{text}_"
                
                para_parts.append(text)
            
            # Join run parts to form paragraph
            para_text = ''.join(para_parts).strip()
            
            if para_text:
                # Detect if this is a list item
                if para_text.startswith(('-', '•', '●', '○')):
                    # Bullet list
                    clean_text = re.sub(r'^[-•●○]\s*', '', para_text)
                    markdown_parts.append(f"- {clean_text}")
                elif re.match(r'^\d+\.\s+', para_text):
                    # Numbered list - keep as is
                    markdown_parts.append(para_text)
                else:
                    # Regular paragraph
                    markdown_parts.append(para_text)
        
        # Join paragraphs with proper markdown line breaks (blank line = new paragraph)
        return '\n\n'.join(markdown_parts)  # Double newline for markdown paragraphs

    
    def detect_email_sections(self) -> List[Dict[str, any]]:
        """
        Detect if document contains multiple email templates.
        Each table represents one email with its own campaign name from the header.
        Returns list of dicts with campaign_name and extracted data.
        """
        tables_data = self.extract_table_content()
        
        # tables_data is already a list of dicts with 'campaign_name' and 'data'
        # Just need to format it properly for the rest of the code
        emails = []
        for table_info in tables_data:
            emails.append({
                'name': table_info['campaign_name'],  # e.g., "Email 1: Long-form email"
                'campaign_name': table_info['campaign_name'],
                'data': table_info['data']
            })
        
        return emails if emails else []
    
    def prepare_row(self, data: Dict[str, str], metadata: Optional[Dict] = None) -> Dict[str, str]:
        """Prepare a single row for Google Sheets"""
        # Apply transformations to data
        data = self.combine_salutation_with_study_info(data)
        data = self.apply_red_field_defaults(data)
        data = self.extract_cta_from_study_info(data)
        
        # Map extracted data to columns
        mapped_data = self.map_to_columns(data)
        
        # Create row with all standard columns
        row = {}
        for col in self.STANDARD_COLUMNS:
            if col in mapped_data:
                # For multi-option fields, use checkbox detection if enabled
                if col in ['Email Subject Line', 'Banner Headline', 'Hero Image (URL)']:
                    row[col] = self.extract_selected_option(mapped_data[col], col)
                else:
                    # Clean up the text
                    text = mapped_data[col]
                    # Normalize line breaks
                    text = text.replace('\r\n', '\n').replace('\r', '\n')
                    # Remove excessive blank lines
                    text = re.sub(r'\n{3,}', '\n\n', text)
                    row[col] = text.strip()
            else:
                row[col] = ''
        
        # Add metadata if provided
        if metadata:
            row.update(metadata)
        
        return row
    
    def extract_all_emails(self, base_metadata: Optional[Dict] = None) -> List[Dict[str, str]]:
        """Extract all email templates from document"""
        emails = self.detect_email_sections()
        rows = []
        
        for idx, email in enumerate(emails, 1):
            metadata = dict(base_metadata) if base_metadata else {}
            
            # Keep FULL Message Name from email header (e.g., "Email 1: Long-form email")
            # Campaign Name comes from user input
            email_header = email.get('name', f'Email {idx}')
            message_name = email_header  # Keep full text
            
            # Set Message Name in metadata (overrides any user-provided value)
            metadata['Message Name'] = message_name
            
            # For warning tracking, use full name
            self.current_message_name = message_name
            
            row = self.prepare_row(email['data'], metadata)
            rows.append(row)
        
        return rows
    
    def get_warnings(self) -> List[Dict[str, str]]:
        """Get list of warnings generated during extraction"""
        return self.warnings
    
    def has_warnings(self) -> bool:
        """Check if any warnings were generated"""
        return len(self.warnings) > 0
    
    def format_warnings(self) -> str:
        """Format warnings as a readable string"""
        if not self.warnings:
            return "No warnings"
        
        lines = ["\n" + "="*80]
        lines.append("WARNING: Missing Checkbox Selections")
        lines.append("="*80)
        lines.append("\nThe following fields have multiple options but no checkbox [X] was found:")
        lines.append("")
        
        # Group warnings by message name
        from collections import defaultdict
        warnings_by_message = defaultdict(list)
        for warning in self.warnings:
            message_name = warning.get('message_name', 'Unknown Email')
            warnings_by_message[message_name].append(warning)
        
        # Display warnings grouped by message
        for message_name, warnings in sorted(warnings_by_message.items()):
            lines.append(f"\n  Email: {message_name}")
            for warning in warnings:
                lines.append(f"     - {warning['field']}")
                lines.append(f"       > {warning['action']}")
        
        lines.append("\n" + "-"*80)
        lines.append("Recommendation: Review these fields in the Word document and add [X] to mark")
        lines.append("the preferred option, then run the extraction again.")
        lines.append("="*80)
        
        return "\n".join(lines)


class GoogleSheetsExporter:
    """Export campaign data for Google Sheets"""
    
    def __init__(self, columns: List[str]):
        self.columns = columns
    
    def to_tsv(self, rows: List[Dict[str, str]], output_file: str) -> str:
        """Export to TSV format (Google Sheets friendly)"""
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=self.columns, delimiter='\t', extrasaction='ignore')
            writer.writeheader()
            writer.writerows(rows)
        return output_file
    
    def to_clipboard(self, rows: List[Dict[str, str]]) -> bool:
        """Copy rows to clipboard for direct paste"""
        if not CLIPBOARD_AVAILABLE:
            return False
        
        lines = []
        for row in rows:
            values = [
                str(row.get(col, '')).replace('\t', '    ').replace('\n', ' ')
                for col in self.columns
            ]
            lines.append('\t'.join(values))
        
        clipboard_text = '\n'.join(lines)
        
        try:
            pyperclip.copy(clipboard_text)
            return True
        except Exception:
            return False

    def to_markdown(self, rows: List[Dict[str, str]], output_file: str) -> str:
        """Export to Markdown format with preserved formatting"""
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("# Campaign Content Export\n\n")
            
            for idx, row in enumerate(rows, 1):
                campaign_name = row.get('Campaign Name', 'Unnamed Campaign')
                message_name = row.get('Message Name', 'Unnamed Message')
                
                f.write(f"## Email {idx}: {message_name}\n\n")
                f.write(f"**Campaign:** {campaign_name}\n\n")
                f.write("---\n\n")
                
                # Write all fields
                for col in self.columns:
                    value = row.get(col, '')
                    if value and col not in ['Campaign Name', 'Message Name']:
                        # Format field name as heading
                        f.write(f"### {col}\n\n")
                        # Write value (already in markdown if formatted)
                        f.write(f"{value}\n\n")
                
                f.write("\n---\n\n")
        
        return output_file

    
    def display_preview(self, rows: List[Dict[str, str]]):
        """Display preview of extracted data"""
        print("\n" + "=" * 100)
        print("EXTRACTED CONTENT PREVIEW")
        print("=" * 100)
        
        for idx, row in enumerate(rows, 1):
            print(f"\n{'-' * 100}")
            print(f"ROW {idx}: {row.get('Message Name', 'Unnamed')}")
            print('-' * 100)
            
            for col in self.columns:
                value = row.get(col, '')
                if value:
                    # Truncate for display
                    display_value = value[:150] + "..." if len(value) > 150 else value
                    display_value = display_value.replace('\n', ' | ')
                    print(f"\n  * {col}")
                    print(f"     {display_value}")


def collect_metadata(smart_defaults: Dict[str, str] = None) -> Dict[str, str]:
    """Interactively collect campaign metadata with smart defaults"""
    print("\n" + "=" * 100)
    print("CAMPAIGN METADATA")
    print("=" * 100)
    print("This information will be added to all extracted emails.")
    print("(Message Name will be auto-extracted from document headers)")
    
    if smart_defaults:
        print("\n[Smart defaults detected from document:]")
        for key, value in smart_defaults.items():
            if key != 'Campaign Name':  # Campaign Name handled separately
                print(f"  {key}: {value}")
    
    metadata = {}
    
    # Campaign Name with smart default
    default_campaign = smart_defaults.get('Campaign Name', '') if smart_defaults else ''
    prompt = f"\n> Campaign Name"
    if default_campaign:
        prompt += f" [{default_campaign}]"
    prompt += ": "
    campaign_name = input(prompt).strip()
    if campaign_name:
        metadata['Campaign Name'] = campaign_name
    elif default_campaign:
        metadata['Campaign Name'] = default_campaign
    
    # Language with smart default
    default_lang = smart_defaults.get('Language', 'en-us') if smart_defaults else 'en-us'
    language = input(f"> Language (en-us or en-es) [{default_lang}]: ").strip()
    metadata['Language'] = language if language else default_lang
    
    # URL/UTM with smart default
    default_url = smart_defaults.get('URL/UTM', '') if smart_defaults else ''
    if default_url:
        url_prompt = f"> URL/UTM [{default_url}]: "
    else:
        url_prompt = "> URL/UTM (leave empty if in document): "
    url = input(url_prompt).strip()
    if url:
        metadata['URL/UTM'] = url
    elif default_url:
        metadata['URL/UTM'] = default_url
    
    # Sponsor with smart default
    default_sponsor = smart_defaults.get('Sponsor', '') if smart_defaults else ''
    if default_sponsor:
        sponsor_prompt = f"> Sponsor Name (for end matter) [{default_sponsor}]: "
    else:
        sponsor_prompt = "> Sponsor Name (for end matter): "
    sponsor = input(sponsor_prompt).strip()
    metadata['End Matter (Enter Sponsor)'] = sponsor if sponsor else (default_sponsor if default_sponsor else '')
    
    return metadata


def main():
    """Main execution function"""
    # Set UTF-8 encoding for Windows console
    import sys
    if sys.platform == 'win32':
        try:
            sys.stdout.reconfigure(encoding='utf-8')
        except:
            pass
    
    print("=" * 100)
    print("WALGREENS CAMPAIGN CONTENT EXTRACTOR")
    print("   Extract Word doc content -> Ready for Google Sheets paste")
    print("=" * 100)
    
    # Find Word documents in JSONGenerator folder
    json_gen_path = Path("JSONGenerator")
    if not json_gen_path.exists():
        json_gen_path = Path(".")
    
    docx_files = list(json_gen_path.glob("*.docx"))
    
    if not docx_files:
        print("\n[!] No .docx files found in JSONGenerator folder")
        docx_path = input("\n> Enter path to Word document: ").strip()
        if not docx_path:
            return
    elif len(docx_files) == 1:
        docx_path = str(docx_files[0])
        print(f"\n[+] Found document: {docx_files[0].name}")
    else:
        print(f"\n[+] Found {len(docx_files)} Word documents:")
        for idx, file in enumerate(docx_files, 1):
            print(f"   {idx}. {file.name}")
        
        choice = input(f"\nSelect document (1-{len(docx_files)}): ").strip()
        try:
            docx_path = str(docx_files[int(choice) - 1])
        except (ValueError, IndexError):
            print("[!] Invalid selection")
            return
    
    print(f"\n[+] Processing: {Path(docx_path).name}")
    
    # Extract smart defaults first
    print("\n" + "=" * 100)
    print("ANALYZING DOCUMENT...")
    print("=" * 100)
    try:
        temp_extractor = CampaignExtractor(docx_path)
        smart_defaults = temp_extractor.extract_smart_defaults_from_document()
        print("\n[+] Smart defaults detected:")
        if smart_defaults.get('Campaign Name'):
            print(f"    Campaign Name: {smart_defaults['Campaign Name']}")
        if smart_defaults.get('Sponsor'):
            print(f"    Sponsor: {smart_defaults['Sponsor']}")
        if smart_defaults.get('Language'):
            print(f"    Language: {smart_defaults['Language']}")
        if smart_defaults.get('URL/UTM'):
            print(f"    URL/UTM: {smart_defaults['URL/UTM']}")
    except Exception as e:
        print(f"\n[!] Could not extract smart defaults: {e}")
        smart_defaults = None
    
    # Ask for metadata
    add_metadata = input("\n[?] Add/confirm campaign metadata? (y/n) [y]: ").strip().lower()
    metadata = collect_metadata(smart_defaults) if add_metadata != 'n' else smart_defaults
    
    # Extract content
    print("\n" + "=" * 100)
    print("EXTRACTING CONTENT...")
    print("=" * 100)
    
    try:
        extractor = CampaignExtractor(docx_path)
        rows = extractor.extract_all_emails(metadata)
        
        print(f"\n[+] Extracted {len(rows)} email template(s)")
        
        # Show warnings if any
        if extractor.has_warnings():
            print(extractor.format_warnings())
        
        # Preview
        exporter = GoogleSheetsExporter(CampaignExtractor.STANDARD_COLUMNS)
        exporter.display_preview(rows)
        
        # Export to TSV
        output_file = Path(docx_path).stem + '_for_google_sheets.tsv'
        tsv_path = exporter.to_tsv(rows, output_file)
        
        print("\n" + "=" * 100)
        print("EXPORT COMPLETE")
        print("=" * 100)
        print(f"\n[+] TSV file created: {output_file}")
        print("   -> Open and copy/paste into Google Sheets")
        
        # Try clipboard
        if exporter.to_clipboard(rows):
            print("\n[+] Content copied to clipboard!")
            print("   -> Paste directly into Google Sheets (Ctrl+V / Cmd+V)")
            print(f"   -> {len(rows)} row(s) ready to paste")
        else:
            if not CLIPBOARD_AVAILABLE:
                print("\n[TIP] Install pyperclip for automatic clipboard copy:")
                print("   pip install pyperclip")
            else:
                print("\n[!] Could not copy to clipboard, use TSV file instead")
        
        # Summary
        print("\n" + "=" * 100)
        print("NEXT STEPS:")
        print("=" * 100)
        print("1. Open your Google Sheets campaign spreadsheet")
        print("2. Click on the first empty row")
        print("3. Press Ctrl+V (Windows) or Cmd+V (Mac)")
        print("4. Review the pasted content")
        print("5. Add Hero Image URLs if needed")
        print("6. Check the box in column A to generate JSON!")
        
        # Save JSON for reference
        json_file = Path(docx_path).stem + '_extracted.json'
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(rows, f, indent=2, ensure_ascii=False)
        print(f"\n[+] Also saved as JSON: {json_file}")
        
        # Save Markdown with formatting
        markdown_file = Path(docx_path).stem + '_formatted.md'
        exporter.to_markdown(rows, markdown_file)
        print(f"[+] Also saved as Markdown: {markdown_file}")
        
    except Exception as e:
        print(f"\n[!] Error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()

