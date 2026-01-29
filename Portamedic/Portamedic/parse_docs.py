#!/usr/bin/env python3
"""
Script to parse all old documentation files into .txt format
and create a clean HTML summary.
"""

import os
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup

# Base paths
OLD_DOCS_PATH = Path("/Users/noahdebrincat/Desktop/Projects/Portamedic/old docs")
NEW_DOCS_PATH = Path("/Users/noahdebrincat/Desktop/Projects/Portamedic/new docs")

# Ensure new docs folder exists
NEW_DOCS_PATH.mkdir(exist_ok=True)

def extract_docx_text(docx_path):
    """Extract text from a .docx file."""
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        return f"Error extracting text: {e}"

def extract_html_text(html_path):
    """Extract text from an HTML file."""
    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator='\n', strip=True)
    except Exception as e:
        return f"Error extracting text: {e}"

# Manually extracted content from images (OCR equivalent)
IMAGE_TEXTS = {
    # api for portamedic folder images
    "api_for_portamedic_1.png": """SOAP 1.1

The following is a sample SOAP 1.1 request and response. The placeholders shown need to be replaced with actual values.

POST /MainService.asmx HTTP/1.1
Host: localhost
Content-Type: text/xml; charset=utf-8
Content-Length: length
SOAPAction: "http://tempuri.org/CheckAvailability"

<?xml version="1.0" encoding="utf-8"?>
<soap:Envelope xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/">
  <soap:Body>
    <CheckAvailability xmlns="http://tempuri.org/">
      <strUserName>string</strUserName>
      <strUserPassword>string</strUserPassword>
      <strAccountNumber>string</strAccountNumber>
      <intAge>short</intAge>
      <intAmount>int</intAmount>
      <strGender>string</strGender>
      <strUniqueID>string</strUniqueID>
      <strZipCode>string</strZipCode>
      <strExceptionProduct>string</strExceptionProduct>
      <strApplicationState>string</strApplicationState>
      <datReservation>dateTime</datReservation>
    </CheckAvailability>
  </soap:Body>
</soap:Envelope>

HTTP/1.1 200 OK
Content-Type: text/xml; charset=utf-8
Content-Length: length

<?xml version="1.0" encoding="utf-8"?>
<soap:Envelope xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/">
  <soap:Body>
    <CheckAvailabilityResponse xmlns="http://tempuri.org/">
      <CheckAvailabilityResult>
        <ArrayOfString>
          <string>string</string>
          <string>string</string>
        </ArrayOfString>
        <ArrayOfString>
          <string>string</string>
          <string>string</string>
        </ArrayOfString>
      </CheckAvailabilityResult>
    </CheckAvailabilityResponse>
  </soap:Body>
</soap:Envelope>""",

    "api_for_portamedic_2.png": """CheckAvailability

Test

To test the operation using the HTTP POST protocol, click the 'Invoke' button.

Parameter | Value
strUserName: ABC
strUserPassword: 123
strAccountNumber: 03034
intAge: 50
intAmount: 500000
strGender: Male
strUniqueID: 123
strZipCode: 99999
strExceptionProduct: (empty)
strApplicationState: NY
datReservation: 12/25/2015

Sample Return XML with Examiner ID's and available times

<?xml version="1.0" encoding="UTF-8"?>
<ArrayOfArrayOfString xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <ArrayOfString>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
    <string>NSC00000000000000000000020820587</string>
  </ArrayOfString>
  <ArrayOfString>
    <string>1:00 PM</string>
    <string>1:15 PM</string>
    <string>1:30 PM</string>
    <string>1:45 PM</string>
    <string>2:00 PM</string>
    <string>2:15 PM</string>
    <string>2:30 PM</string>
    <string>2:45 PM</string>
    <string>3:00 PM</string>
    <string>3:15 PM</string>
    <string>3:30 PM</string>
  </ArrayOfString>
</ArrayOfArrayOfString>""",

    "api_for_portamedic_3.png": """Make Reservation

The carrier will pass the same parameters they used to check availability; however, they will also include the selected Examiner ID and available time selected as illustrated below.

SOAP 1.1

The following is a sample SOAP 1.1 request and response. The placeholders shown need to be replaced with actual values.

POST /MainService.asmx HTTP/1.1
Host: localhost
Content-Type: text/xml; charset=utf-8
Content-Length: length
SOAPAction: "http://tempuri.org/MakeReservation"

<?xml version="1.0" encoding="utf-8"?>
<soap:Envelope xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/">
  <soap:Body>
    <MakeReservation xmlns="http://tempuri.org/">
      <strUserName>string</strUserName>
      <strUserPassword>string</strUserPassword>
      <strAccountNumber>string</strAccountNumber>
      <intAge>short</intAge>
      <intAmount>int</intAmount>
      <strGender>string</strGender>
      <strUniqueID>string</strUniqueID>
      <strZipCode>string</strZipCode>
      <strExceptionProduct>string</strExceptionProduct>
      <strApplicationState>string</strApplicationState>
      <datReservation>dateTime</datReservation>
      <strExaminerID>string</strExaminerID>
      <strTime>string</strTime>
      <strSpecialInstruction>string</strSpecialInstruction>
    </MakeReservation>
  </soap:Body>
</soap:Envelope>

HTTP/1.1 200 OK
Content-Type: text/xml; charset=utf-8
Content-Length: length

<?xml version="1.0" encoding="utf-8"?>
<soap:Envelope xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/">
  <soap:Body>
    <MakeReservationResponse xmlns="http://tempuri.org/">
      <MakeReservationResult>
        <IsSuccessful>boolean</IsSuccessful>
        <ConfirmationNumber>string</ConfirmationNumber>
        <ErrorMessage>string</ErrorMessage>
        <ClientID>string</ClientID>
      </MakeReservationResult>
    </MakeReservationResponse>
  </soap:Body>
</soap:Envelope>""",

    "api_for_portamedic_4.png": """WebService1

Click here for a complete list of operations.

MakeReservation

Test

To test the operation using the HTTP POST protocol, click the 'Invoke' button.

Parameter | Value
strUserName: ABC
strUserPassword: 123
strAccountNumber: 03034
intAge: 50
intAmount: 500000
strGender: Male
strUniqueID: 123
strZipCode: 99999
strExceptionProduct: (empty)
strApplicationState: NY
datReservation: 12/25/2015
strExaminerID: NSC00000000000000000000020820587
strTime: 1:00 PM
strSpecialInstruction: This is a test.

Successful Return XML Sample

<?xml version="1.0" encoding="UTF-8"?>
<ResponseObject xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <IsSuccessful>true</IsSuccessful>
  <ConfirmationNumber>bea5cf96-dc25-4390-adfa-76e6cb2973a5</ConfirmationNumber>
  <ErrorMessage/>
  <ClientID>123</ClientID>
</ResponseObject>

Unsuccessful Return XML Sample – Message is time/examiner selected is no longer available

<?xml version="1.0" encoding="UTF-8"?>
<ResponseObject xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <IsSuccessful>false</IsSuccessful>
  <ConfirmationNumber/>
  <ErrorMessage>Specified time is not available for this examiner - 1:00 PM - WRONGINFOHERE---------- :o)</ErrorMessage>
  <ClientID>123</ClientID>
</ResponseObject>""",

    "api_for_portamedic_5.png": """Covering Office Zip Code

The carrier will have the ability to pass APPS the zip code of the applicant. In return, APPS will provide the local field office contact information that will cover the zip codes exam.

SOAP 1.1

The following is a sample SOAP 1.1 request and response. The placeholders shown need to be replaced with actual values.

POST /MainService.asmx HTTP/1.1
Host: localhost
Content-Type: text/xml; charset=utf-8
Content-Length: length
SOAPAction: "http://tempuri.org/OfficeCoveringZipCode"

<?xml version="1.0" encoding="utf-8"?>
<soap:Envelope xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/">
  <soap:Body>
    <OfficeCoveringZipCode xmlns="http://tempuri.org/">
      <strUserName>string</strUserName>
      <strUserPassword>string</strUserPassword>
      <strZipCode>string</strZipCode>
    </OfficeCoveringZipCode>
  </soap:Body>
</soap:Envelope>

HTTP/1.1 200 OK
Content-Type: text/xml; charset=utf-8
Content-Length: length

<?xml version="1.0" encoding="utf-8"?>
<soap:Envelope xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/">
  <soap:Body>
    <OfficeCoveringZipCodeResponse xmlns="http://tempuri.org/">
      <OfficeCoveringZipCodeResult>
        <OfficeNumber>string</OfficeNumber>
        <DisplayName>string</DisplayName>
        <Street>string</Street>
        <AptNumber>string</AptNumber>
        <City>string</City>
        <State>string</State>
        <ZipCode>string</ZipCode>
        <Phone>string</Phone>
        <Fax>string</Fax>
        <EmailAlias>string</EmailAlias>
      </OfficeCoveringZipCodeResult>
    </OfficeCoveringZipCodeResponse>
  </soap:Body>
</soap:Envelope>""",

    "api_for_portamedic_6.png": """WebService1

Click here for a complete list of operations.

OfficeCoveringZipCode

Test

To test the operation using the HTTP POST protocol, click the 'Invoke' button.

Parameter | Value
strUserName: ABC
strUserPassword: 123
strZipCode: 11364

Sample Return XML

<?xml version="1.0" encoding="UTF-8"?>
<OfficeInfo xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <OfficeNumber>01100</OfficeNumber>
  <DisplayName>New York City, New York</DisplayName>
  <Street>42-40 Bell Blvd.</Street>
  <AptNumber>305</AptNumber>
  <City>Bayside</City>
  <State>NY</State>
  <ZipCode>11361</ZipCode>
  <Phone>718-225-9300</Phone>
  <Fax>718-423-9505</Fax>
  <EmailAlias>NewYorkCity.NY@APPSdom.com</EmailAlias>
</OfficeInfo>""",
}

# Technical connectivity folder images
TECHNICAL_IMAGE_TEXTS = {
    "image1.png": "TECHNICAL CONNECTIVITY - Schedule4Real API (Title page)",
    "image2.png": "TECHNICAL CONNECTIVITY | Schedule4Real API (Header)",
    "image3.png": "(Decorative blue gradient header)",
    "image4.png": """SOAP 1.1

The following is a sample SOAP 1.1 request and response. The placeholders shown need to be replaced with actual values.

POST /MainService.asmx HTTP/1.1
Host: localhost
Content-Type: text/xml; charset=utf-8
Content-Length: length
SOAPAction: "http://tempuri.org/CheckAvailability"

CheckAvailability Request Parameters:
- strUserName: string
- strUserPassword: string
- strAccountNumber: string
- intAge: short
- intAmount: int
- strGender: string
- strUniqueID: string
- strZipCode: string
- strExceptionProduct: string
- strApplicationState: string
- datReservation: dateTime

CheckAvailability Response:
Returns ArrayOfString with examiner IDs and available times.""",

    "image5.png": """Sample Return XML with Examiner ID's and available times

<?xml version="1.0" encoding="UTF-8"?>
<ArrayOfArrayOfString xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <ArrayOfString>
    <string>NSC00000000000000000000020820587</string>
    (11 examiner IDs)
  </ArrayOfString>
  <ArrayOfString>
    <string>1:00 PM</string>
    <string>1:15 PM</string>
    <string>1:30 PM</string>
    <string>1:45 PM</string>
    <string>2:00 PM</string>
    <string>2:15 PM</string>
    <string>2:30 PM</string>
    <string>2:45 PM</string>
    <string>3:00 PM</string>
    <string>3:15 PM</string>
    <string>3:30 PM</string>
  </ArrayOfString>
</ArrayOfArrayOfString>""",

    "image6.png": """WebService1 - CheckAvailability Test

Parameter | Value
strUserName: ABC
strUserPassword: 123
strAccountNumber: 03034
intAge: 50
intAmount: 500000
strGender: Male
strUniqueID: 123
strZipCode: 99999
strExceptionProduct: (empty)
strApplicationState: NY
datReservation: 12/25/2015""",

    "image7.png": """WebService1 - MakeReservation Test

Parameter | Value
strUserName: ABC
strUserPassword: 123
strAccountNumber: 03034
intAge: 50
intAmount: 500000
strGender: Male
strUniqueID: 123
strZipCode: 99999
strExceptionProduct: (empty)
strApplicationState: NY
datReservation: 12/25/2015
strExaminerID: NSC00000000000000000000020820587
strTime: 1:00 PM
strSpecialInstruction: This is a test.""",

    "image8.png": """SOAP 1.1 - MakeReservation

The following is a sample SOAP 1.1 request and response. The placeholders shown need to be replaced with actual values.

MakeReservation Request Parameters (same as CheckAvailability plus):
- strExaminerID: string
- strTime: string
- strSpecialInstruction: string

MakeReservation Response:
- IsSuccessful: boolean
- ConfirmationNumber: string
- ErrorMessage: string
- ClientID: string""",

    "image9.png": """Unsuccessful Return XML Sample

<?xml version="1.0" encoding="UTF-8"?>
<ResponseObject xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <IsSuccessful>false</IsSuccessful>
  <ConfirmationNumber/>
  <ErrorMessage>Specified time is not available for this examiner - 1:00 PM - WRONGINFOHERE---------- :o)</ErrorMessage>
  <ClientID>123</ClientID>
</ResponseObject>""",

    "image10.png": """Successful Return XML Sample

<?xml version="1.0" encoding="UTF-8"?>
<ResponseObject xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <IsSuccessful>true</IsSuccessful>
  <ConfirmationNumber>bea5cf96-dc25-4390-adfa-76e6cb2973a5</ConfirmationNumber>
  <ErrorMessage/>
  <ClientID>123</ClientID>
</ResponseObject>""",

    "image11.png": """WebService1 - OfficeCoveringZipCode Test

Parameter | Value
strUserName: ABC
strUserPassword: 123
strZipCode: 11364""",

    "image12.png": "(Tim McCabe signature/name)",

    "image13.png": """SOAP 1.1 - OfficeCoveringZipCode

Request Parameters:
- strUserName: string
- strUserPassword: string
- strZipCode: string

Response:
- OfficeNumber: string
- DisplayName: string
- Street: string
- AptNumber: string
- City: string
- State: string
- ZipCode: string
- Phone: string
- Fax: string
- EmailAlias: string""",

    "image14.png": """API Error Messages List

Error | Description
"strUserName cannot be blank" | Username is required
"strUserPassword cannot be blank" | Password is required
"The user is invalid" | Username/password combo not in db
"strAccountNumber must be 5 chars long." | Account number must be 5 chars
"strAccountNumberCompanion must be 5 chars long." | Account number must be 5 chars
"intAge is invalid, < 1 or > 120." | Age is invalid
"intAgeCompanion is invalid, < 1 or > 120." | Age is invalid
"intAmount is invalid, < 0 or > 50000000" | Insurance amount is invalid
"intAmountCompanion is invalid, < 0 or > 50000000" | Insurance amount is invalid
"strGender is invalid, (M|F|U|Male|Female|Unknown|{blank})" | Gender is invalid
"strGenderCompanion is invalid, (M|F|U|Male|Female|Unknown|{blank})" | Gender is invalid
"strUniqueID cannot be empty." | Unique ID is required
"strUniqueIDCompanion cannot be empty." | Unique ID is required
"strUniqueID already exists in the database." | Unique ID must be unique/already existing in db
"strUniqueIDCompanion already exists in the database." | Unique ID must be unique/already existing in db
"strZipCode must be at least 5 digits long" | Zip code must be at least 5 chars long
"strZipCode first five chars must contain numerics only" | Zip code must be numeric
"datReservation must be at least 5 days out" | Reservation date is too soon
"AccountNumber is invalid" | Account number not found in db
"AccountNumberCompanion is invalid" | Account number not found in db
"An Application State is needed." | Requirements vary by state for this account, app state needed
"An Exception Product is needed." | Requirements vary by product for this account, product needed
"No requirements are needed." | No requirements found for APPS to perform
"There are no examiners who can cover this scenario." | No examiners cover this zip code with required equipment
"There are no examiners who can cover this date." | Examiners cover this zip code with required equipment, just not on this date""",

    "image15.png": "APPS Para Medical Services Logo",

    "image16.png": """Sample Return XML - OfficeCoveringZipCode

<?xml version="1.0" encoding="UTF-8"?>
<OfficeInfo xmlns="http://tempuri.org/" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <OfficeNumber>01100</OfficeNumber>
  <DisplayName>New York City, New York</DisplayName>
  <Street>42-40 Bell Blvd.</Street>
  <AptNumber>305</AptNumber>
  <City>Bayside</City>
  <State>NY</State>
  <ZipCode>11361</ZipCode>
  <Phone>718-225-9300</Phone>
  <Fax>718-423-9505</Fax>
  <EmailAlias>NewYorkCity.NY@APPSdom.com</EmailAlias>
</OfficeInfo>""",

    "image17.png": "(Blank/decorative image)",
}


def main():
    print("=" * 60)
    print("PORTAMEDIC OLD DOCS PARSER")
    print("=" * 60)
    
    all_content = []
    
    # 1. Parse DOCX files in api for portamedic folder
    api_folder = OLD_DOCS_PATH / "api for portamedic"
    docx_files = list(api_folder.glob("*.docx"))
    
    print(f"\n📄 Processing {len(docx_files)} DOCX files from 'api for portamedic'...")
    
    for docx_file in docx_files:
        print(f"   - {docx_file.name}")
        text = extract_docx_text(docx_file)
        output_file = NEW_DOCS_PATH / f"docx_api_portamedic_{docx_file.stem}.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"SOURCE: {docx_file.name}\n")
            f.write("=" * 60 + "\n\n")
            f.write(text)
        all_content.append(("DOCX: " + docx_file.name, text))
    
    # 2. Parse main Portamedic HTML file (extracted from DOCX)
    main_html = OLD_DOCS_PATH / "PM - Ordering API Documentation.docx" / "PMOrderingAPIDocumentation.docx.html"
    print(f"\n📄 Processing main Portamedic documentation (HTML)...")
    if main_html.exists():
        text = extract_html_text(main_html)
        output_file = NEW_DOCS_PATH / "portamedic_ordering_api_v2.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"SOURCE: {main_html.name}\n")
            f.write("=" * 60 + "\n\n")
            f.write(text)
        all_content.append(("HTML: PM Ordering API Documentation", text))
        print(f"   ✅ Created portamedic_ordering_api_v2.txt")
    else:
        print(f"   ⚠️ File not found: {main_html}")
    
    # 3. Parse HTML file
    html_file = OLD_DOCS_PATH / "Technical connectivity - Schedule4Real API 2.docx" / "TechnicalconnectivitySchedule4RealAPI.docx.html"
    if html_file.exists():
        print(f"\n📄 Processing Schedule4Real HTML documentation...")
        text = extract_html_text(html_file)
        output_file = NEW_DOCS_PATH / "schedule4real_api.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"SOURCE: {html_file.name}\n")
            f.write("=" * 60 + "\n\n")
            f.write(text)
        all_content.append(("HTML: Schedule4Real API", text))
    
    # 4. Save image text extractions - api for portamedic
    print(f"\n🖼️ Processing {len(IMAGE_TEXTS)} images from 'api for portamedic'...")
    combined_api_images = []
    for img_name, text in IMAGE_TEXTS.items():
        print(f"   - {img_name}")
        output_file = NEW_DOCS_PATH / f"{img_name.replace('.png', '')}.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"SOURCE: {img_name}\n")
            f.write("=" * 60 + "\n\n")
            f.write(text)
        combined_api_images.append(text)
        all_content.append(("IMAGE: " + img_name, text))
    
    # 5. Save image text extractions - Technical connectivity
    print(f"\n🖼️ Processing {len(TECHNICAL_IMAGE_TEXTS)} images from 'Technical connectivity'...")
    combined_tech_images = []
    for img_name, text in TECHNICAL_IMAGE_TEXTS.items():
        print(f"   - {img_name}")
        output_file = NEW_DOCS_PATH / f"technical_{img_name.replace('.png', '')}.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"SOURCE: Technical connectivity - {img_name}\n")
            f.write("=" * 60 + "\n\n")
            f.write(text)
        combined_tech_images.append(text)
        all_content.append(("IMAGE: Technical - " + img_name, text))
    
    # 6. Create combined text file with all content
    print("\n📋 Creating combined documentation file...")
    combined_file = NEW_DOCS_PATH / "ALL_DOCUMENTATION_COMBINED.txt"
    with open(combined_file, 'w', encoding='utf-8') as f:
        f.write("PORTAMEDIC API DOCUMENTATION - COMPLETE COMPILATION\n")
        f.write("=" * 70 + "\n")
        f.write("Generated from old docs folder\n")
        f.write("=" * 70 + "\n\n")
        
        for source, text in all_content:
            f.write(f"\n{'='*70}\n")
            f.write(f"SOURCE: {source}\n")
            f.write(f"{'='*70}\n\n")
            f.write(text)
            f.write("\n\n")
    
    print(f"\n✅ All text files created in: {NEW_DOCS_PATH}")
    
    # 7. Create HTML Summary
    create_html_summary(all_content)
    
    print("\n" + "=" * 60)
    print("DONE! All files processed successfully.")
    print("=" * 60)


def create_html_summary(all_content):
    """Create a clean 1-page HTML summary of all documentation."""
    print("\n📝 Creating HTML summary...")
    
    html_content = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Portamedic API Documentation Summary</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;600&family=Outfit:wght@300;400;600;700&display=swap');
        
        :root {
            --bg-primary: #0a0e14;
            --bg-secondary: #12171f;
            --bg-card: #151b24;
            --accent-primary: #00d9ff;
            --accent-secondary: #6366f1;
            --accent-success: #10b981;
            --accent-warning: #f59e0b;
            --accent-error: #ef4444;
            --text-primary: #e8eaed;
            --text-secondary: #9ca3af;
            --text-muted: #6b7280;
            --border-color: #1e2530;
        }
        
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Outfit', sans-serif;
            background: var(--bg-primary);
            color: var(--text-primary);
            line-height: 1.7;
            min-height: 100vh;
            background-image: 
                radial-gradient(ellipse at 20% 0%, rgba(0, 217, 255, 0.08) 0%, transparent 50%),
                radial-gradient(ellipse at 80% 100%, rgba(99, 102, 241, 0.08) 0%, transparent 50%);
        }
        
        .container {
            max-width: 1100px;
            margin: 0 auto;
            padding: 3rem 2rem;
        }
        
        header {
            text-align: center;
            margin-bottom: 3rem;
            padding: 2rem;
            background: linear-gradient(135deg, var(--bg-secondary) 0%, var(--bg-card) 100%);
            border-radius: 16px;
            border: 1px solid var(--border-color);
            position: relative;
            overflow: hidden;
        }
        
        header::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 3px;
            background: linear-gradient(90deg, var(--accent-primary), var(--accent-secondary));
        }
        
        .logo {
            font-size: 0.9rem;
            letter-spacing: 4px;
            text-transform: uppercase;
            color: var(--accent-primary);
            margin-bottom: 0.5rem;
            font-weight: 600;
        }
        
        h1 {
            font-size: 2.5rem;
            font-weight: 700;
            background: linear-gradient(135deg, var(--text-primary) 0%, var(--accent-primary) 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            margin-bottom: 0.5rem;
        }
        
        .subtitle {
            color: var(--text-secondary);
            font-size: 1.1rem;
            font-weight: 300;
        }
        
        .grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(320px, 1fr));
            gap: 1.5rem;
            margin-bottom: 2rem;
        }
        
        .card {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 1.5rem;
            transition: all 0.3s ease;
        }
        
        .card:hover {
            border-color: var(--accent-primary);
            transform: translateY(-2px);
            box-shadow: 0 8px 30px rgba(0, 217, 255, 0.1);
        }
        
        .card-header {
            display: flex;
            align-items: center;
            gap: 0.75rem;
            margin-bottom: 1rem;
        }
        
        .card-icon {
            width: 40px;
            height: 40px;
            border-radius: 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.2rem;
        }
        
        .card-icon.blue { background: rgba(0, 217, 255, 0.15); }
        .card-icon.purple { background: rgba(99, 102, 241, 0.15); }
        .card-icon.green { background: rgba(16, 185, 129, 0.15); }
        .card-icon.orange { background: rgba(245, 158, 11, 0.15); }
        
        .card h3 {
            font-size: 1.1rem;
            font-weight: 600;
            color: var(--text-primary);
        }
        
        .card p {
            color: var(--text-secondary);
            font-size: 0.95rem;
        }
        
        .section {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 2rem;
            margin-bottom: 2rem;
        }
        
        .section h2 {
            font-size: 1.4rem;
            margin-bottom: 1.5rem;
            color: var(--accent-primary);
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }
        
        .endpoint {
            background: var(--bg-secondary);
            border-radius: 8px;
            padding: 1rem;
            margin-bottom: 1rem;
            border-left: 3px solid var(--accent-primary);
        }
        
        .endpoint-name {
            font-family: 'JetBrains Mono', monospace;
            font-size: 1rem;
            color: var(--accent-primary);
            font-weight: 600;
        }
        
        .endpoint-desc {
            color: var(--text-secondary);
            font-size: 0.9rem;
            margin-top: 0.5rem;
        }
        
        .params-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(180px, 1fr));
            gap: 0.5rem;
            margin-top: 0.75rem;
        }
        
        .param {
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.8rem;
            padding: 0.25rem 0.5rem;
            background: var(--bg-primary);
            border-radius: 4px;
            color: var(--text-muted);
        }
        
        .urls-box {
            background: var(--bg-secondary);
            border-radius: 8px;
            padding: 1rem 1.5rem;
            margin-top: 1rem;
        }
        
        .url-item {
            display: flex;
            align-items: center;
            gap: 0.75rem;
            margin: 0.5rem 0;
        }
        
        .url-label {
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 1px;
            color: var(--text-muted);
            min-width: 80px;
        }
        
        .url-value {
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.85rem;
            color: var(--accent-success);
        }
        
        .error-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 0.85rem;
        }
        
        .error-table th,
        .error-table td {
            padding: 0.75rem 1rem;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }
        
        .error-table th {
            background: var(--bg-secondary);
            color: var(--accent-primary);
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 1px;
            font-size: 0.75rem;
        }
        
        .error-table td:first-child {
            font-family: 'JetBrains Mono', monospace;
            color: var(--accent-error);
            font-size: 0.8rem;
        }
        
        footer {
            text-align: center;
            padding: 2rem;
            color: var(--text-muted);
            font-size: 0.85rem;
        }
        
        .badge {
            display: inline-block;
            padding: 0.25rem 0.75rem;
            border-radius: 20px;
            font-size: 0.75rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }
        
        .badge-post { background: rgba(0, 217, 255, 0.2); color: var(--accent-primary); }
        .badge-soap { background: rgba(99, 102, 241, 0.2); color: var(--accent-secondary); }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <div class="logo">APPS Para Medical Services</div>
            <h1>Portamedic API</h1>
            <p class="subtitle">Schedule4Real Integration Documentation</p>
        </header>
        
        <div class="grid">
            <div class="card">
                <div class="card-header">
                    <div class="card-icon blue">📋</div>
                    <h3>Scheduling API</h3>
                </div>
                <p>SOAP 1.1 Web Service for checking examiner availability and making reservations for paramedical exams.</p>
            </div>
            
            <div class="card">
                <div class="card-header">
                    <div class="card-icon purple">🔐</div>
                    <h3>Authentication</h3>
                </div>
                <p>Username and password credentials required for all API requests. Account numbers map to carrier/distribution channels.</p>
            </div>
            
            <div class="card">
                <div class="card-header">
                    <div class="card-icon green">🏥</div>
                    <h3>Office Lookup</h3>
                </div>
                <p>Find local field office contact information based on applicant's zip code for exam coverage.</p>
            </div>
            
            <div class="card">
                <div class="card-header">
                    <div class="card-icon orange">⚡</div>
                    <h3>Real-time</h3>
                </div>
                <p>Live availability checking with immediate confirmation or error messages for reservation conflicts.</p>
            </div>
        </div>
        
        <section class="section">
            <h2>🌐 API Endpoints</h2>
            
            <div class="urls-box">
                <div class="url-item">
                    <span class="url-label">Test</span>
                    <span class="url-value">https://www.integratedtestingservices.com/S4R/mainservice.asmx</span>
                </div>
                <div class="url-item">
                    <span class="url-label">Production</span>
                    <span class="url-value">https://appslive.com/S4R/mainservice.asmx</span>
                </div>
            </div>
            
            <div style="margin-top: 1.5rem;">
                <div class="endpoint">
                    <div style="display: flex; align-items: center; gap: 0.75rem;">
                        <span class="badge badge-soap">SOAP</span>
                        <span class="endpoint-name">CheckAvailability</span>
                    </div>
                    <p class="endpoint-desc">Returns available examiners and time slots for a given zip code and date.</p>
                    <div class="params-grid">
                        <span class="param">strUserName</span>
                        <span class="param">strUserPassword</span>
                        <span class="param">strAccountNumber</span>
                        <span class="param">intAge</span>
                        <span class="param">intAmount</span>
                        <span class="param">strGender</span>
                        <span class="param">strUniqueID</span>
                        <span class="param">strZipCode</span>
                        <span class="param">strExceptionProduct</span>
                        <span class="param">strApplicationState</span>
                        <span class="param">datReservation</span>
                    </div>
                </div>
                
                <div class="endpoint">
                    <div style="display: flex; align-items: center; gap: 0.75rem;">
                        <span class="badge badge-soap">SOAP</span>
                        <span class="endpoint-name">MakeReservation</span>
                    </div>
                    <p class="endpoint-desc">Books an appointment with a specific examiner at a selected time slot.</p>
                    <div class="params-grid">
                        <span class="param">+ strExaminerID</span>
                        <span class="param">+ strTime</span>
                        <span class="param">+ strSpecialInstruction</span>
                    </div>
                </div>
                
                <div class="endpoint">
                    <div style="display: flex; align-items: center; gap: 0.75rem;">
                        <span class="badge badge-soap">SOAP</span>
                        <span class="endpoint-name">OfficeCoveringZipCode</span>
                    </div>
                    <p class="endpoint-desc">Returns local field office contact info for a given zip code.</p>
                    <div class="params-grid">
                        <span class="param">strUserName</span>
                        <span class="param">strUserPassword</span>
                        <span class="param">strZipCode</span>
                    </div>
                </div>
            </div>
        </section>
        
        <section class="section">
            <h2>📝 Required Parameters</h2>
            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); gap: 1rem;">
                <div class="endpoint">
                    <span class="endpoint-name">Zip Code</span>
                    <p class="endpoint-desc">5-character code</p>
                </div>
                <div class="endpoint">
                    <span class="endpoint-name">Gender</span>
                    <p class="endpoint-desc">M/F (or Male/Female)</p>
                </div>
                <div class="endpoint">
                    <span class="endpoint-name">Unique ID</span>
                    <p class="endpoint-desc">OrderNo or Policy#</p>
                </div>
                <div class="endpoint">
                    <span class="endpoint-name">Account #</span>
                    <p class="endpoint-desc">CompanyID (5 digits)</p>
                </div>
                <div class="endpoint">
                    <span class="endpoint-name">Requirements</span>
                    <p class="endpoint-desc">BLOOD, EKG, URINE, VITALS</p>
                </div>
                <div class="endpoint">
                    <span class="endpoint-name">Reservation Date</span>
                    <p class="endpoint-desc">MM/DD/YYYY (5+ days out)</p>
                </div>
            </div>
        </section>
        
        <section class="section">
            <h2>⚠️ Common Error Messages</h2>
            <div style="overflow-x: auto;">
                <table class="error-table">
                    <thead>
                        <tr>
                            <th>Error Code</th>
                            <th>Description</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>strUserName cannot be blank</td>
                            <td>Username is required</td>
                        </tr>
                        <tr>
                            <td>strUserPassword cannot be blank</td>
                            <td>Password is required</td>
                        </tr>
                        <tr>
                            <td>The user is invalid</td>
                            <td>Username/password combo not found</td>
                        </tr>
                        <tr>
                            <td>strAccountNumber must be 5 chars long</td>
                            <td>Account number format invalid</td>
                        </tr>
                        <tr>
                            <td>intAge is invalid</td>
                            <td>Age must be between 1 and 120</td>
                        </tr>
                        <tr>
                            <td>intAmount is invalid</td>
                            <td>Insurance amount must be 0-50,000,000</td>
                        </tr>
                        <tr>
                            <td>strZipCode must be at least 5 digits</td>
                            <td>Zip code format invalid</td>
                        </tr>
                        <tr>
                            <td>datReservation must be at least 5 days out</td>
                            <td>Reservation date is too soon</td>
                        </tr>
                        <tr>
                            <td>No examiners who can cover this scenario</td>
                            <td>No coverage for zip code with required equipment</td>
                        </tr>
                    </tbody>
                </table>
            </div>
        </section>
        
        <footer>
            <p>Generated from Portamedic Old Documentation • Schedule4Real API v2</p>
            <p style="margin-top: 0.5rem; color: var(--text-muted);">APPS Para Medical Services</p>
        </footer>
    </div>
</body>
</html>'''
    
    output_file = NEW_DOCS_PATH / "documentation_summary.html"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print(f"   ✅ HTML summary created: {output_file}")


if __name__ == "__main__":
    main()

